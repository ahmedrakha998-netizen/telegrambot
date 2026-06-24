import nest_asyncio
nest_asyncio.apply()

import os
import logging
import json
import asyncio
import re  # 💡 تم إضافة مكتبة regex للتصفية الذكية
from docx import Document
from pypdf import PdfReader
from google import genai
from google.genai import types
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, CallbackQueryHandler

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)


# 🔒 قراءة التوكنات بشكل آمن من متغيرات البيئة بالسيرفر دون كشفها في الكود
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

# تأكد من أن السيرفر يحتوي على المفاتيح فعلاً قبل تشغيل البوت
if not TELEGRAM_TOKEN or not GEMINI_API_KEY:
    raise ValueError("❌ خطأ: لم يتم العثور على TELEGRAM_TOKEN أو GEMINI_API_KEY في متغيرات بيئة السيرفر!")

# إعداد عميل Gemini
ai_client = genai.Client(api_key=GEMINI_API_KEY)


def extract_text_from_file(file_path):
    text = ""
    file_name = file_path.lower()
    
    if file_name.endswith('.docx'):
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip(): text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): text += cell.text + "\n"
                    
    elif file_name.endswith('.pdf'):
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted: 
                    text += extracted + "\n"
        except Exception as pdf_err:
            logging.error(f"Error reading PDF pages: {pdf_err}")
            
    # 🔍 فحص ما إذا كان النص يحتوي على حروف عربية قبل حذفها
    has_arabic = bool(re.search(r'[\u0600-\u06FF]', text))
    
    # 🌟 الفلتر الذكي: الإبقاء فقط على الحروف الإنجليزية، الأرقام، وعلامات الترقيم والمسافات
    english_only_text = re.sub(r'[^a-zA-Z0-9\s.,;:?!()\'"\[\]\-\+\*/\=]', '', text)
    
    return english_only_text.strip(), has_arabic

def generate_questions_with_ai(content_text, mode="ai", is_continuation=False):
    if mode == "file":
        instruction = "The text provided IS ALREADY a list of questions (e.g., an exam or a Q&A bank). Extract ALL of them exactly as they are without skipping any or creating new ones. Ignore any introductory text or non-question sentences."
    else:
        if is_continuation:
            instruction = "This is a continuation request. Generate a NEW set of 20-25 questions from different sections of the text, avoiding the questions you generated previously."
        else:
            instruction = "If the text is a lecture, textbook chapter, or summary, generate a comprehensive quiz consisting of around 25 concise multiple-choice questions (MCQs) covering all sections."

    prompt = f"""
    You are an expert quiz generator. Analyze the following text.
    {instruction}

    Each question must have between 2 to 4 options. One option must be the correct answer.

    CRITICAL: Return the response ONLY as a valid JSON array, containing objects with this exact structure:
    [
      {{
        "question": "The text of the question",
        "options": ["Option 1", "Option 2", "Option 3", "Option 4"],
        "correct_index": 0
      }}
    ]
    Note: correct_index is a 0-based integer pointing to the correct answer in the options list.
    Language: Generate the questions in the SAME language as the exam questions in the input text (English).

    Text to analyze:
    {content_text}
    """

    try:
        response = ai_client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json"
            )
        )
        response_text = response.text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
        return json.loads(response_text)
    except Exception as e:
        logging.error(f"Gemini AI Error: {e}")
        return []

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.chat_data['mode'] = 'ai'
    await update.message.reply_text(
        "  احلى مسا يبويا اهلا بيك !\n\n"
        " يمكنك التحكم في طريقة عمل البوت عبر الأوامر:\n"
        "🤖 /ai : لوضع الذكاء الاصطناعي (ترسل محاضرة والـ AI يعمل أسئلة من عنده).\n"
        "📝 /file : لوضع ملف الأسئلة (تبعت ملف يحتوي على أسئلة جاهزة ليحولها لكويز مباشرة).\n\n"
        "النظام النشط حالياً: 🤖 وضع الذكاء الاصطناعي."
        "\n\nاو ابعت الملف علي طول!"
    )

async def set_mode_ai(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.chat_data['mode'] = 'ai'
    await update.message.reply_text("🤖 تم تفعيل [وضع الذكاء الاصطناعي]. أرسل المحاضرة أو الملخص الآن.")

async def set_mode_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.chat_data['mode'] = 'file'
    await update.message.reply_text("📝 تم تفعيل [وضع ملف الأسئلة]. أرسل الملف الذي يحتوي على الأسئلة الجاهزة الآن.")

async def send_quiz_polls(chat_id, questions, bot, show_continue=True):
    for q in questions:
        question_text = q['question'][:300]
        options = [str(opt)[:100] for opt in q['options'] if str(opt).strip()]
        correct_id = int(q['correct_index'])

        if correct_id >= len(options):
            correct_id = 0

        if len(options) >= 2:
            try:
                await bot.send_poll(
                    chat_id=chat_id,
                    question=question_text,
                    options=options,
                    type='quiz',
                    correct_option_id=correct_id,
                    is_anonymous=False
                )
                await asyncio.sleep(1.8)
            except Exception as poll_err:
                logging.error(f"Error sending poll: {poll_err}")

    if show_continue:
        keyboard = [
            [
                InlineKeyboardButton("نعم، أضف المزيد ➕", callback_data="continue_yes"),
                InlineKeyboardButton("لا، شطبنا ❌", callback_data="continue_no")
            ]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await bot.send_message(
            chat_id=chat_id,
            text=" خلصنا شوية اسالة تحب نكمل من نفس الملف ؟",
            reply_markup=reply_markup
        )
    else:
        await bot.send_message(chat_id=chat_id, text=" تم استخراج وتحويل جميع الأسئلة من ملفك بنجاح! شطبنا كده.")

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    document = update.message.document
    file_name = document.file_name.lower()
    current_mode = context.chat_data.get('mode', 'ai')

    if not (file_name.endswith('.docx') or file_name.endswith('.pdf')):
        await update.message.reply_text("عذراً، أدعم ملفات Word و PDF فقط.")
        return

    status_msg = "📥 جاري معالجة المحاضرة بالـ AI... (الملفات الكبيرة قد تستغرق دقيقة، يرجى الانتظار) ⏳" if current_mode == 'ai' else "📥 جاري قراءة واستخراج الأسئلة من ملفك... (يرجى الانتظار قليلاً) ⏳"
    await update.message.reply_text(status_msg)

    _, file_extension = os.path.splitext(file_name)
    temp_file_name = f"temp_quiz_file{file_extension}"
    local_path = os.path.join(os.getcwd(), temp_file_name)

    tg_file = await context.bot.get_file(document.file_id)
    await tg_file.download_to_drive(local_path)

    try:
        # 💡 استلام النص مع علامة فحص وجود نصوص عربية
        raw_text, contains_arabic = extract_text_from_file(local_path)
        if not raw_text:
            await update.message.reply_text("الملف فارغ أو لا يحتوي على نصوص إنجليزية قابلة للقراءة.")
            return

        # ⚠️ تنبيه تلقائي للمستخدم إذا وجدنا نصوصاً جانبية بالعربية
        if contains_arabic:
            await update.message.reply_text("⚠️ تنبيه: تم رصد نصوص  باللغة العربية داخل الملف،سيتم تجاهلها وارسال الاسالة في دقائق.")

        context.chat_data['last_file_text'] = raw_text
        questions = generate_questions_with_ai(raw_text, mode=current_mode, is_continuation=False)

        if not questions:
            await update.message.reply_text("واجه الـ AI مشكلة في معالجة الملف وصياغة الأسئلة. تأكد من صلاحية الـ Gemini API Key الخاص بك.")
            return

        await update.message.reply_text(f"✅ تم معالجة {len(questions)} سؤال بنجاح! جاري الإرسال الآن...")
        show_buttons = True if current_mode == 'ai' else False
        await send_quiz_polls(update.effective_chat.id, questions, context.bot, show_continue=show_buttons)

    except Exception as e:
        logging.error(f"Error: {e}")
        await update.message.reply_text(" حدث خطأ أثناء معالجة الملف الرجاء الرجوع لطارق لحل المشكلة .")
    finally:
        if os.path.exists(local_path):
            os.remove(local_path)

async def handle_callback_query(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    if query.data == "continue_yes":
        raw_text = context.chat_data.get('last_file_text')
        current_mode = context.chat_data.get('mode', 'ai')
        if not raw_text:
            await query.edit_message_text("❌ عذراً، انتهت صلاحية الملف المؤقتة. يرجى إعادة إرسال الملف مجدداً.")
            return

        await query.edit_message_text(" تم استقبال طلبك! جاري مراجعة الملف وتوليد دفعة جديدة من الأسئلة...")
        questions = generate_questions_with_ai(raw_text, mode=current_mode, is_continuation=True)

        if not questions:
            await context.bot.send_message(chat_id=update.effective_chat.id, text="لم يجد الـ AI أفكاراً جديدة لصياغة أسئلة أخرى من هذا المستند.")
            return

        await context.bot.send_message(chat_id=update.effective_chat.id, text=f" رائع! تم توليد {len(questions)} سؤال جديد، جاري إرسالها...")
        await send_quiz_polls(update.effective_chat.id, questions, context.bot, show_continue=True)

    elif query.data == "continue_no":
        await query.edit_message_text("شطبنا كده يبويا بالتوفيق يا نجم")

def main():
    application = Application.builder().token(TELEGRAM_TOKEN).build()
    
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("ai", set_mode_ai))
    application.add_handler(CommandHandler("file", set_mode_file))
    
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    application.add_handler(CallbackQueryHandler(handle_callback_query))
    
    print("احلى مسا واحلى مزاج")
    application.run_polling()

if __name__ == '__main__':
    main()
